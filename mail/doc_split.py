import docx



def print_tables(doc):
  """Prints docx fyle content to the console

  Args:
      doc (.docx): .docx file
  """
  for table in doc.tables:
    for row in table.rows:
      for cell in row.cells:
        for paragraph in cell.paragraphs:
          print(paragraph.text)                    
        print_tables(cell)
        
            
def iter_tables_within_table(table):
  """Functions used to find nested tables within table

  Args:
      table (docx table): docx document object table

  Yields:
      docx table: nested table object
  """
  for row in table.rows:
    for cell in row.cells:
      for nested_table in cell.tables:
        if nested_table.cell(5,1).text.split()[0] == 'Darbuotojas:':
          yield nested_table
          yield from iter_tables_within_table(nested_table)    


def delete_row_in_table(row : int, file_name: str, employees: list):
  """Deletes all rows besides specified row

  Args:
      row (int): row index
      file_name (string): .docx file name
      employees (list): list of employees information
  """
  document = docx.Document(file_name)
  vowels = {ord('ą') : 'a', ord('č'): 'c', ord('ę'): 'e', ord('ė'): 'e', ord('į'): 'i', ord('š'): 's', ord('ų'):'u', ord('ū'): 'u', ord('ž'): 'z'}
  if len(document.tables[0].rows) > row:        
    for i in range(0,row):
      document.tables[0]._tbl.remove(document.tables[0].rows[0]._tr)
    for i in range(0, len(document.tables[0].rows) - 1):
      document.tables[0]._tbl.remove(document.tables[0].rows[1]._tr)
        
    string = f'{employees[row][1]}_{employees[row][2]}_atlyginimo_lapelis.docx'
    employees[row].append('dranseika.mantas@gmail.com')
    # employees[row].append(f'{employees[row][1]}.{employees[row][2]}@matom.ai'.lower().translate(vowels))
    # employees[row].append(f'{employees[row][1]}@matom.ai'.lower().translate(vowels))
    employees[row].append(string)       
    document.save(string)    


def split_file(file_name: str) -> list:
  """splits salary slips .docx file into individual slips and returns employees data

  Args:
      file_name (string): .docx file name

  Returns:
      list: employees list
  """
  document = docx.Document(file_name)
  table = document.tables[0]
  
  employees = []
  for nested_table in iter_tables_within_table(table):
    employees.append(nested_table.cell(5,1).text.split())
      
  for i in range(0, len(document.tables[0].rows)):
    delete_row_in_table(i, file_name, employees)
      
  return employees
        
# print(split_file('atlyginimo_lapeliai.docx'))
